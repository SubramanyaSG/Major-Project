# src/utils/text_processing.py - Complete Text Processing

import re
import os
from pathlib import Path

def clean_text(text):
    """Remove special characters and clean text thoroughly."""
    
    # Remove URLs
    text = re.sub(r'http\S+|www\S+|https\S+', '', text, flags=re.MULTILINE)
    
    # Remove email addresses
    text = re.sub(r'\S+@\S+', '', text)
    
    # Remove extra whitespace and newlines
    text = re.sub(r'\s+', ' ', text)
    
    # Keep only alphanumeric, spaces, and basic punctuation
    text = re.sub(r'[^a-zA-Z0-9\s.,!?\'"-]', '', text)
    
    # Remove multiple punctuation
    text = re.sub(r'[.,!?]{2,}', '.', text)
    
    # Clean up spaces around punctuation
    text = re.sub(r'\s+([.,!?])', r'\1', text)
    
    # Remove extra whitespace again
    text = ' '.join(text.split())
    
    return text.strip()


def extract_text_from_file(file_path):
    """
    Extract text from .txt, .docx, or .pdf files.
    
    Args:
        file_path: Path to the file
    
    Returns:
        Extracted text string
    """
    
    file_path = str(file_path)
    file_ext = file_path.lower().split('.')[-1]
    
    print(f"📄 Extracting text from: {os.path.basename(file_path)}")
    print(f"   File type: {file_ext}")
    
    try:
        if file_ext == 'txt':
            return _extract_from_txt(file_path)
        
        elif file_ext == 'docx':
            return _extract_from_docx(file_path)
        
        elif file_ext == 'pdf':
            return _extract_from_pdf(file_path)
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    except Exception as e:
        print(f"❌ Error extracting text: {e}")
        raise Exception(f"Failed to extract text: {str(e)}")


def _extract_from_txt(file_path):
    """Extract text from TXT file."""
    
    print("  🔍 Reading TXT file...")
    
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        text = None
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"  ✅ Successfully read with encoding: {encoding}")
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise Exception("Could not decode file with any encoding")
        
        # Clean the text
        text = clean_text(text)
        
        print(f"  ✅ Extracted {len(text)} characters")
        
        return text
    
    except Exception as e:
        print(f"  ❌ Error: {e}")
        raise


def _extract_from_docx(file_path):
    """Extract text from DOCX file."""
    
    print("  🔍 Reading DOCX file...")
    
    try:
        from docx import Document
        
        # Open document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        # Combine all text
        text = ' '.join(text_parts)
        
        # Clean the text
        text = clean_text(text)
        
        print(f"  ✅ Extracted {len(text)} characters from DOCX")
        
        return text
    
    except ImportError:
        print("  ❌ python-docx not installed")
        print("     Run: pip install python-docx")
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    
    except Exception as e:
        print(f"  ❌ Error: {e}")
        raise


def _extract_from_pdf(file_path):
    """Extract text from PDF file."""
    
    print("  🔍 Reading PDF file...")
    
    try:
        import PyPDF2
        
        text_parts = []
        
        # Open and read PDF
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            
            print(f"  📄 PDF has {num_pages} pages")
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                    
                    if (page_num + 1) % 5 == 0:
                        print(f"  Processing page {page_num + 1}/{num_pages}")
                
                except Exception as e:
                    print(f"  ⚠️ Warning: Could not extract from page {page_num + 1}: {e}")
                    continue
        
        # Combine all text
        text = ' '.join(text_parts)
        
        # Clean the text
        text = clean_text(text)
        
        print(f"  ✅ Extracted {len(text)} characters from PDF")
        
        return text
    
    except ImportError:
        print("  ❌ PyPDF2 not installed")
        print("     Run: pip install PyPDF2")
        raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
    
    except Exception as e:
        print(f"  ❌ Error: {e}")
        raise


def validate_extracted_text(text, min_length=10):
    """Validate that extracted text is usable."""
    
    if not text or len(text) < min_length:
        return False, f"Text too short (min {min_length} characters)"
    
    if len(text) > 10000:
        print(f"⚠️ Text length {len(text)} exceeds 10000, truncating...")
        return True, text[:10000]
    
    return True, text
